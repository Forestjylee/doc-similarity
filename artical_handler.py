from docx import Document

'''
读取.docx文件
@:param  artical_text
@:param  artical_path
by: Junyi
'''


class ArticalHandler(object):

    def __init__(self, artical_text=None, artical_path=None):
        self.artical_path = artical_path
        self.artical_text = artical_text

    def get_words(self):
        if self.artical_text:
            for item in self.artical_text:
                yield item
        elif self.artical_path:
            for item in self.get_words_from_docx(docx_path=self.artical_path):
                yield item

    @staticmethod
    def get_words_from_docx(docx_path, include_table=True):
        document = Document(docx_path)
        for paragraph in document.paragraphs:
            if paragraph.text:
                yield paragraph.text
        if include_table:
            tables = document.tables
            for table in tables:
                for row in range(0, len(table.rows)):
                    for column in range(0, len(table.columns)):
                        yield table.cell(row, column).text
