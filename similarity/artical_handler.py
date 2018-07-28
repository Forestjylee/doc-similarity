'''
读取.docx文件
@:param  artical_directory
by: Junyi
'''
from docx import Document
import redis
import os

# TODO 实现get_data_from_antiword(部署时)使用的默认读入方式

class ArticalHandler(object):

    def __init__(self, artical_directory):
        self.artical_directory = artical_directory

    # 返回文章内容迭代器
    def get_artical_generators(self):
        artical_paths = os.listdir(self.artical_directory)
        for artical_path in artical_paths:
            docx_path = os.path.join(self.artical_directory, artical_path)
            artical_generator = self.get_words_from_docx(docx_path)
            if artical_generator:
                yield artical_generator

    @staticmethod
    def get_words_from_docx(docx_path, include_table=True):
        try:
            document = Document(docx_path)
        except:
            print("{}不是docx类型的文档".format(docx_path))
            yield None
        for paragraph in document.paragraphs:
            if paragraph.text:
                yield paragraph.text
        if include_table:
            tables = document.tables
            for table in tables:
                for row in range(0, len(table.rows)):
                    for column in range(0, len(table.columns)):
                        yield table.cell(row, column).text

    def get_word_from_antiword(self):
        pass